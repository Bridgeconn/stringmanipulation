# Converting table docx files to normal .docx or .txt format. 


from docx import Document
import glob
import os
import re

# Checking the docx directory and accessing it.
try:
	if not os.path.exists("File with table"):
		print("No folder for File with table, Create a folder with name \"File with table\" and drop docx file in that")
	else:
		os.chdir("File with table")
except OSError:
	print ('Error: Creating directory.' + "File with table")

files = glob.glob("*.docx")

for docx_file in files:
	# reading the contents from docx file
	document = Document(docx_file)
	table = document.tables[0]
	content = ""

	# Fetching the content row-wise from the table
	for doc in table.rows:
		try:
			for paragraph in doc.cells[0].paragraphs:
				content += paragraph.text + "\n"
		except:
			content += "" + "\n"

	os.chdir("..")

	# Creating a folder or accessing a folder for converted sbv files
	try:
		if not os.path.exists("docx"):
			os.mkdir("docx")
			os.chdir("docx")
		else:
			os.chdir("docx")
	except OSError:
		print ('Error: Creating directory.' + "docx")
	
	doc = Document()

	# Creating new file if not available
	if not docx_file:
		doc.save(docx_file)
	
	doc.add_paragraph(content)
	doc.save(docx_file)
	print(content)
	
	# For .txt file, I don't need so commented
	#file_name = re.sub(r".docx","",docx_file)
	#txt_file = open(file_name + ".txt", 'w')
	#txt_file.write(content)
	#print(content)
	
	os.chdir("..")