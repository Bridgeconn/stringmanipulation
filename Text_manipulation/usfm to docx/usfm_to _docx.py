'''
 Script to convert usfm file to docx by removing tags and unwanted lines.
'''


import glob
import os
import re
from docx import Document
import docx

'''The name of the folder which contains the usfm files'''
folder = "usfm"

for root, dirs, files in os.walk(folder):
	for file in files:
		os.chdir("usfm")
		cwd = os.getcwd()
		print(file,cwd)

		'''Reads the usfm file'''
		usfm_file = open(file, 'r')
		content = usfm_file.read()
		main_content = ""
		docx_file = ""
		lines = content.split("\n")

		'''Change directory'''
		os.chdir("..")	
		doc = Document()
		
		try:
			if not os.path.exists("docx"):
				os.mkdir("docx")
				os.chdir("docx")
			else:
				os.chdir("docx")
		except OSError:
			print ('Error: Creating directory.' + "docx")

		for line in lines:
			if line:
				main_content = ""

				'''Regex to remove tags and unwanted lines'''
				search_book = re.match(r"(\\id)\s(\w+)",line)
				search_chapter = re.match(r"(\\c)\s(\d+)",line)
				search_sub = re.match(r"(\\s)\s?(.*)",line)
				search_verse = re.match(r"(\\v)\s?(.*)",line)
				search_quotes = re.match(r"(\\q\d?)\s?(.*)",line)

				if search_book:
					file_name = search_book.group(2)
					main_content = file_name
					print(main_content)

					'''Opens a new docx file with same file name'''
					docx_file = file_name + ".docx"
					doc.add_paragraph(main_content)
					doc.save(docx_file)
				
				elif search_chapter:
					chapter = search_chapter.group(2)
					main_content = "\n" + "अध्याय " + chapter
					print(main_content)
					doc.add_paragraph(main_content)
					doc.save(docx_file)

				
				elif search_sub:
					sub = search_sub.group(2)
					main_content = "\n" + "# " + sub + " #"
					print(main_content)
					
					doc.add_paragraph(main_content)
					#pa = doc.add_paragraph()
					#pa.add_run(main_content).bold = True
				
					doc.save(docx_file)


				elif search_verse:
					verse = search_verse.group(2)
					main_content = verse
					print(main_content)
					doc.add_paragraph(main_content)
					doc.save(docx_file)

				elif search_quotes:
					quote = search_quotes.group(2)
					main_content = quote
					print(main_content)
					doc.add_paragraph(main_content)
					doc.save(docx_file)

		os.chdir("..")