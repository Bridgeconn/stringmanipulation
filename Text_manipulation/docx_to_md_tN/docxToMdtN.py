'''
The script is written for creating a md file with the content of docx file.
File conversion from docx to md.
#!/usr/bin/python
# -*- coding: utf-8 -*-
'''

from docx import Document
import os
import re
import sys
import glob
import errno

files = glob.glob('**/*.docx')
for name in files:
    filename = os.path.splitext(name)[0]
    print (filename)

    ''' Initialization of variables '''
    document = Document(name)
    table = document.tables[0]
    row_count = len(table.rows)
    pre_folder = ""
    chunknum = ""
    pre_chunk0 = ""
    pre_chunk1 = ""
    a = 1
    x = 2
    #prev_to_prev_chunk = ""
    pattern = re.compile(r'-')
    f = filename.split('.')
    folder = f[0]

    ''' Check if directory exists '''
    if not os.path.exists(folder):
        os.mkdir(folder)

    ''' Convert into md file '''
    print ("Conversion in progress !")
    content = ""
    for r in range(2, row_count):

        ''' Read table from docx file '''
        row = table.rows[r]
        cell = row.cells[2]
        paragraph = cell.paragraphs[0]
        p = paragraph.text
        
        if p:

            ''' Check for chunk numbers and create md files'''
            if (str(p[-1]).isdigit()):
                x = 1
                a = 0
                
                if pattern.search(str(p)):
                   
                    chunknum = str(p).split('-')
                    #folder = chunknum[0]
                    if content:
                        
                        #print(content)
                        print(pre_folder+"/"+pre_chunk0+"/"+pre_chunk1+".md")
                        outfile = open(pre_folder+"/"+pre_chunk0+"/"+pre_chunk1+".md", "a")
                        outfile.write(content)
                        #cwd = ""
                        content = ""
                        print(chunknum)
                        if not os.path.exists('%s/%s' %(folder, chunknum[0])):
                            os.mkdir("%s/%s/" %(folder, chunknum[0]))
                            #cwd = os.getcwd()
                            outfile = open('%s/%s/%s.md' %(folder, chunknum[0],chunknum[1]), "w")
                            pre_folder = folder
                            pre_chunk0 = chunknum[0]
                            pre_chunk1 = chunknum[1]
                        else:
                            outfile = open('%s/%s/%s.md' %(folder, chunknum[0],chunknum[1]), "w")
                            pre_folder = folder
                            pre_chunk0 = chunknum[0]
                            pre_chunk1 = chunknum[1]
                        #x = 1
                    else:
                        if not os.path.exists('%s/%s' %(folder, chunknum[0])):
                            os.mkdir("%s/%s/" %(folder, chunknum[0]))
                            outfile = open('%s/%s/%s.md' %(folder, chunknum[0],chunknum[1]), "w")
                            pre_folder = folder
                            pre_chunk0 = chunknum[0]
                            pre_chunk1 = chunknum[1]
                        else:
                            pre_folder = folder
                            pre_chunk0 = chunknum[0]
                            pre_chunk1 = chunknum[1]
                        

            ''' Check for answers '''
            elif p:
                if a ==0:
                    search_brace = re.findall(r"\(",p)
                    if search_brace:
                        x = 1
                    a = 1

                if x == 0:
                    content += p + "\n"
                    x = 1
                    print(p + "\n")
                elif x == 1:
                    if search_brace:
                        content += p + "\n"
                        x = 1
                        search_brace = ""
                        print(p + "\n")
                    else:
                        content += "# " + p + "\n"
                        x = 0
                        print("# " + p + "\n")
