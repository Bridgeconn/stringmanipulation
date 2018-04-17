# Docx content to docx with table for translation


from docx import Document
import glob
import os

files = glob.glob("*.docx")
for file in files:
	document = Document(file)
	try:
		if not os.path.exists("Docx with table"):
			os.makedirs("Docx with table")
			os.chdir("Docx with table")
		else:
			os.chdir("Docx with table")
	except OSError:
		print ('Error: Creating directory.' + "Docx with table")

	content = ""
	row = 0

	table_document = Document()

	# Creating new file if not available
	if not file:
		table_document.save(file)
	
	for para in document.paragraphs:
		row += 1

	table = table_document.add_table(rows=row + 1, cols=1)
	
	# For simple table design
	table.style = "TableGrid"
	i = 0

	for doc in document.paragraphs:
		row = table.rows[i]
		row.cells[0].text = doc.text
		print(doc.text)
		i += 1

	# Adding extra row in the end for translation
	row = table.rows[i]
	row.cells[0].text = ""
	table_document.save(file)
	os.chdir("..")
